"""
Shared docx building blocks: cover page, TOC, revision history, helpers.
Uses python-docx library.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement


# ===== Color constants =====
NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_DARK = RGBColor(0x80, 0x80, 0x80)
GRAY_LIGHT = "E7E6E6"
GRAY_HEADER = "F2F2F2"
NAVY_HEX = "1F3864"


# ===== Helpers: cell shading & borders =====
# OOXML CT_TcPrInner schema requires children in this order (relevant subset):
#   cnfStyle, tcW, gridSpan, hMerge, vMerge, tcBorders, shd, noWrap,
#   tcMar, textDirection, tcFitText, vAlign, hideMark
# We therefore remove existing nodes of the same tag and insert at the right spot.

# Order index for tcPr child elements
_TC_ORDER = [
    "cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge",
    "tcBorders", "shd", "noWrap", "tcMar", "textDirection",
    "tcFitText", "vAlign", "hideMark",
]


def _insert_into_tcPr(tcPr, new_elem, tag_name: str):
    """tcPr에 자식 요소를 스키마 순서에 맞게 삽입(기존 동일 태그는 제거)"""
    # 같은 태그 이미 있으면 제거
    for existing in tcPr.findall(qn(f"w:{tag_name}")):
        tcPr.remove(existing)

    # 자기보다 뒤에 와야 할 첫 형제 찾기
    own_idx = _TC_ORDER.index(tag_name)
    insert_before = None
    for child in list(tcPr):
        # child.tag 는 "{namespace}localname" 형식
        local = child.tag.split("}", 1)[-1]
        if local in _TC_ORDER and _TC_ORDER.index(local) > own_idx:
            insert_before = child
            break

    if insert_before is not None:
        insert_before.addprevious(new_elem)
    else:
        tcPr.append(new_elem)


def set_cell_shading(cell, hex_color: str):
    """셀 배경색 설정"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    _insert_into_tcPr(tcPr, shd, "shd")


def set_cell_borders(cell, color="BFBFBF", size="4"):
    """셀 테두리 설정"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    _insert_into_tcPr(tcPr, tcBorders, "tcBorders")


def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    """셀 안쪽 여백"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tcMar.append(node)
    _insert_into_tcPr(tcPr, tcMar, "tcMar")


# OOXML CT_PPrBase schema order (relevant subset):
#   pStyle, keepNext, keepLines, pageBreakBefore, framePr, widowControl,
#   numPr, suppressLineNumbers, pBdr, shd, tabs, suppressAutoHyphens,
#   kinsoku, wordWrap, overflowPunct, topLinePunct, autoSpaceDE, autoSpaceDN,
#   bidi, adjustRightInd, snapToGrid, spacing, ind, contextualSpacing,
#   mirrorIndents, suppressOverlap, jc, textDirection, ...
_PPR_ORDER = [
    "pStyle", "keepNext", "keepLines", "pageBreakBefore", "framePr",
    "widowControl", "numPr", "suppressLineNumbers", "pBdr", "shd",
    "tabs", "suppressAutoHyphens", "kinsoku", "wordWrap", "overflowPunct",
    "topLinePunct", "autoSpaceDE", "autoSpaceDN", "bidi", "adjustRightInd",
    "snapToGrid", "spacing", "ind", "contextualSpacing", "mirrorIndents",
    "suppressOverlap", "jc", "textDirection", "textAlignment", "outlineLvl",
    "rPr",
]


def _insert_into_pPr(pPr, new_elem, tag_name: str):
    """pPr에 자식 요소를 스키마 순서에 맞게 삽입(기존 동일 태그는 제거)"""
    for existing in pPr.findall(qn(f"w:{tag_name}")):
        pPr.remove(existing)

    own_idx = _PPR_ORDER.index(tag_name)
    insert_before = None
    for child in list(pPr):
        local = child.tag.split("}", 1)[-1]
        if local in _PPR_ORDER and _PPR_ORDER.index(local) > own_idx:
            insert_before = child
            break

    if insert_before is not None:
        insert_before.addprevious(new_elem)
    else:
        pPr.append(new_elem)


def add_paragraph_border_top(paragraph, color="1F3864", size="6"):
    """단락 위쪽 테두리 (footer 등에 사용)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), size)
    top.set(qn("w:color"), color)
    pBdr.append(top)
    _insert_into_pPr(pPr, pBdr, "pBdr")


def add_paragraph_border_bottom(paragraph, color="1F3864", size="12"):
    """단락 아래쪽 테두리 (구분선용)"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), size)
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    _insert_into_pPr(pPr, pBdr, "pBdr")


# ===== Document setup =====
def init_document() -> Document:
    """기본 스타일 적용된 새 Document 생성"""
    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Heading styles
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Arial"
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = NAVY

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Arial"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = BLUE

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(11)
    h3.font.bold = True

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # python-docx의 기본 settings.xml에 있는 <w:zoom>이 'percent' 속성 없이
    # 들어가 있어 OOXML 검증에서 오류를 내므로 percent='100' 을 추가해 둔다.
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is not None and zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")

    return doc


# ===== Paragraph helpers =====
def add_paragraph(doc, text: str, bold: bool = False, italic: bool = False,
                  size: int = None, color: RGBColor = None,
                  alignment=None, after_pt: int = 6):
    """일반 단락 추가"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = "Arial"
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after_pt)
    return p


def add_bullet(doc, text: str):
    """불릿 항목 추가"""
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet_bold(doc, bold_part: str, rest: str):
    """앞부분 굵게 + 나머지 일반 텍스트 불릿"""
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_part)
    r1.font.name = "Arial"
    r1.font.size = Pt(11)
    r1.bold = True
    r2 = p.add_run(rest)
    r2.font.name = "Arial"
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_h1(doc, text: str):
    """Heading 1"""
    return doc.add_heading(text, level=1)


def add_h2(doc, text: str):
    """Heading 2"""
    return doc.add_heading(text, level=2)


def add_h3(doc, text: str):
    """Heading 3"""
    return doc.add_heading(text, level=3)


def add_page_break(doc):
    """페이지 나누기"""
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_spacer(doc):
    """빈 줄"""
    doc.add_paragraph()


# ===== Tables =====
def add_info_table(doc, rows: list):
    """2열 라벨-값 표 (label은 회색 배경)"""
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(11.0)

    for i, (label, value) in enumerate(rows):
        # Label cell
        cell_l = table.cell(i, 0)
        cell_l.width = Cm(5.5)
        set_cell_shading(cell_l, GRAY_LIGHT)
        set_cell_borders(cell_l)
        set_cell_margins(cell_l)
        p_l = cell_l.paragraphs[0]
        run_l = p_l.add_run(label)
        run_l.font.name = "Arial"
        run_l.font.size = Pt(11)
        run_l.bold = True

        # Value cell
        cell_v = table.cell(i, 1)
        cell_v.width = Cm(11.0)
        set_cell_borders(cell_v)
        set_cell_margins(cell_v)
        p_v = cell_v.paragraphs[0]
        run_v = p_v.add_run(str(value or ""))
        run_v.font.name = "Arial"
        run_v.font.size = Pt(11)

    return table


def add_before_after_table(doc, rows: list):
    """Item / Before / After 3열 표"""
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.autofit = False
    for col in table.columns:
        col.width = Cm(5.5)

    headers = ["Item", "Before", "After"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.width = Cm(5.5)
        set_cell_shading(cell, NAVY_HEX)
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = WHITE

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Cm(5.5)
            if j == 0:
                set_cell_shading(cell, GRAY_HEADER)
            set_cell_borders(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val or ""))
            run.font.name = "Arial"
            run.font.size = Pt(11)
            if j == 0:
                run.bold = True

    return table


def add_two_col_table(doc, header: list, rows: list, col_widths_cm=(10.5, 6.5)):
    """범용 2열 헤더 표"""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(col_widths_cm[0])
    table.columns[1].width = Cm(col_widths_cm[1])

    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.width = Cm(col_widths_cm[j])
        set_cell_shading(cell, NAVY_HEX)
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = WHITE

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Cm(col_widths_cm[j])
            set_cell_borders(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            run = p.add_run(str(val or ""))
            run.font.name = "Arial"
            run.font.size = Pt(11)

    return table


# ===== Cover Page =====
def add_cover_page(doc, *, doc_title: str, doc_subtitle: str, doc_number: str,
                   rev_no: str, effective_date: str,
                   prepared_by: str = "", reviewed_by: str = "", approved_by: str = ""):
    """표지 페이지 추가"""
    # Top spacer
    for _ in range(3):
        doc.add_paragraph()

    # Top label
    add_paragraph(doc, "MEDICAL DEVICE REGULATORY DOCUMENTATION",
                  bold=True, size=12, color=GRAY_DARK,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=10)

    # Top divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(p, color=NAVY_HEX, size="12")

    # Title
    add_paragraph(doc, doc_title, bold=True, size=22, color=NAVY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=10)

    # Subtitle (multiline)
    for line in doc_subtitle.split("\n"):
        add_paragraph(doc, line, italic=True, size=12, color=BLUE,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=4)

    # Bottom divider line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph_border_bottom(p, color=NAVY_HEX, size="12")

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Document info table (centered)
    info_rows = [
        ("Document Title", doc_title),
        ("Document Number", doc_number),
        ("Revision No.", rev_no),
        ("Effective Date", effective_date),
    ]
    table = doc.add_table(rows=len(info_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(5.0)
    table.columns[1].width = Cm(7.5)

    for i, (label, value) in enumerate(info_rows):
        cell_l = table.cell(i, 0)
        cell_l.width = Cm(5.0)
        set_cell_shading(cell_l, NAVY_HEX)
        set_cell_borders(cell_l)
        set_cell_margins(cell_l, top=140, bottom=140, left=160, right=160)
        p_l = cell_l.paragraphs[0]
        run_l = p_l.add_run(label)
        run_l.font.name = "Arial"
        run_l.font.size = Pt(11)
        run_l.bold = True
        run_l.font.color.rgb = WHITE

        cell_v = table.cell(i, 1)
        cell_v.width = Cm(7.5)
        set_cell_borders(cell_v)
        set_cell_margins(cell_v, top=140, bottom=140, left=160, right=160)
        p_v = cell_v.paragraphs[0]
        run_v = p_v.add_run(str(value))
        run_v.font.name = "Arial"
        run_v.font.size = Pt(11)

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Approval label
    add_paragraph(doc, "DOCUMENT APPROVAL", bold=True, size=13, color=NAVY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=10)

    # Approval table (4 cols: Role / Name / Signature / Date)
    approval_rows = [
        ("Prepared by", prepared_by),
        ("Reviewed by", reviewed_by),
        ("Approved by", approved_by),
    ]
    table = doc.add_table(rows=len(approval_rows) + 1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_widths = [Cm(3.5), Cm(4.5), Cm(4.5), Cm(4.0)]
    for j, w in enumerate(col_widths):
        table.columns[j].width = w

    headers = ["Role", "Name / Title", "Signature", "Date"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.width = col_widths[j]
        set_cell_shading(cell, NAVY_HEX)
        set_cell_borders(cell, color=NAVY_HEX, size="6")
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = WHITE

    for i, (role, name) in enumerate(approval_rows, start=1):
        for j in range(4):
            cell = table.cell(i, j)
            cell.width = col_widths[j]
            if j == 0:
                set_cell_shading(cell, GRAY_HEADER)
            set_cell_borders(cell, color=NAVY_HEX, size="6")
            set_cell_margins(cell, top=300, bottom=300, left=140, right=140)
            p = cell.paragraphs[0]
            text = ""
            bold = False
            if j == 0:
                text = role
                bold = True
            elif j == 1:
                text = name or ""
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.bold = bold

    # Page break to next page
    add_page_break(doc)


# ===== Revision History =====
def add_revision_history(doc, metadata: dict):
    add_h1(doc, "Revision History")

    rows = [
        (metadata.get("revisionNo", "00"),
         metadata.get("effectiveDate", "[YYYY-MM-DD]"),
         "Initial issue.",
         metadata.get("preparedBy", "[Name]")),
        ("", "", "", ""),  # 빈 행
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.autofit = False
    col_widths = [Cm(2.0), Cm(3.0), Cm(7.5), Cm(4.0)]
    for j, w in enumerate(col_widths):
        table.columns[j].width = w

    headers = ["Rev. No.", "Date", "Description of Change", "Author"]
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.width = col_widths[j]
        set_cell_shading(cell, NAVY_HEX)
        set_cell_borders(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = WHITE

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.width = col_widths[j]
            set_cell_borders(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            if j in (0, 1):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(11)

    add_page_break(doc)


# ===== Table of Contents =====
def add_toc(doc):
    add_h1(doc, "Table of Contents")
    add_paragraph(
        doc,
        'Note: After opening this document, right-click the table of contents below and select "Update Field" to refresh page numbers.',
        italic=True, size=9, color=GRAY_DARK, after_pt=6
    )
    add_spacer(doc)

    # TOC field
    p = doc.add_paragraph()
    run = p.add_run()

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click here and select 'Update Field' to populate the table of contents."
    run._r.append(placeholder)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)

    add_page_break(doc)


# ===== Footer =====
def add_footer(doc, doc_number: str):
    """모든 페이지 하단에 문서번호 + 페이지 번호"""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Top border on footer
    add_paragraph_border_top(p, color=NAVY_HEX, size="6")

    # "DOC-NUMBER  |  Page "
    run1 = p.add_run(f"{doc_number}    |    Page ")
    run1.font.name = "Arial"
    run1.font.size = Pt(9)
    run1.font.color.rgb = GRAY_DARK

    # PAGE field
    run_page = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_page._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    run_page._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_page._r.append(fld_end)
    run_page.font.name = "Arial"
    run_page.font.size = Pt(9)
    run_page.font.color.rgb = GRAY_DARK

    # " of "
    run2 = p.add_run(" of ")
    run2.font.name = "Arial"
    run2.font.size = Pt(9)
    run2.font.color.rgb = GRAY_DARK

    # NUMPAGES field
    run_total = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_total._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.text = "NUMPAGES"
    run_total._r.append(instr)
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_total._r.append(fld_end)
    run_total.font.name = "Arial"
    run_total.font.size = Pt(9)
    run_total.font.color.rgb = GRAY_DARK
