"""
FDA 510(k) change review document builder (Python).
Follows the "Regulatory Change Assessment" template shown in FDA's own guidance,
Appendix B: Documentation (Example 2 — complex change with full checklist attached).
"""
from doc_common import (
    init_document, add_paragraph, add_spacer, add_page_break,
    add_before_after_table, add_signature_table, add_labeled_paragraph,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH

FDA_GUIDANCE_TITLE = "Deciding When to Submit a 510(k) for a Change to an Existing Device"

CHECK_YES = "[X]"  # 특정 환경에서 ☒/☐ 유니코드 글리프가 깨지는 것을 피하기 위해 ASCII로 표기
CHECK_NO = "[ ]"

# Which MAIN-chart gate leads into which lettered chart / chart title.
CHART_LETTER = {"MAIN2": "A", "MAIN3": "B", "MAIN4": "C"}
CHART_TITLE = {
    "A": "Labeling Changes",
    "B": "Technology, Engineering, and Performance Changes",
    "C": "Materials Changes",
    "D": "Technology, Engineering, Performance, and Materials Changes for In Vitro Diagnostic Devices",
}


def _checkbox_line(doc, checked: bool, label: str, after_pt: int = 2):
    box = CHECK_YES if checked else CHECK_NO
    add_paragraph(doc, f"{box} {label}", after_pt=after_pt)


def _qa_checkboxes(doc, text_en: str, answer):
    add_paragraph(doc, text_en, after_pt=2)
    _checkbox_line(doc, bool(answer), "Yes", after_pt=0)
    _checkbox_line(doc, not answer, "No", after_pt=10)


def build_fda_document(product_info: dict, change_info: dict,
                       assessment: dict, metadata: dict):
    """FDA 510(k) 변경 검토 문서 생성 — FDA Appendix B 'Regulatory Change Assessment' 양식"""
    doc = init_document()

    # ===== Title =====
    add_paragraph(doc, "Regulatory Change Assessment", bold=True, size=16,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, after_pt=16)

    # ===== Header fields =====
    add_labeled_paragraph(doc, "Product Name", product_info.get("modelName", "[Model Name]"))
    add_labeled_paragraph(doc, "Date of Assessment", metadata.get("effectiveDate", "[YYYY-MM-DD]"))

    device_desc_bits = [b for b in [product_info.get("deviceClass"), product_info.get("manufacturer")] if b]
    device_desc = product_info.get("modelName", "[Model Name]")
    if device_desc_bits:
        device_desc += f" ({', '.join(device_desc_bits)})"
    add_labeled_paragraph(doc, "Device Description", device_desc)

    change_desc = change_info.get("description") or change_info.get("changeTitle", "")
    add_labeled_paragraph(doc, "Description of Change(s)", change_desc)

    add_labeled_paragraph(doc, "Reason for Change(s)", change_info.get("reason", "[Reason for change]"))

    reg_history = f"Most recently cleared 510(k): {product_info.get('fdaK510', 'N/A')}."
    add_labeled_paragraph(
        doc,
        "Applicable Regulatory History (including 510(k) #s and comparison of modified device to most recently cleared version)",
        reg_history,
        after_pt=4,
    )
    add_before_after_table(doc, [(
        change_info.get("componentName", ""),
        change_info.get("beforeValue", ""),
        change_info.get("afterValue", ""),
    )])
    add_spacer(doc)

    # ===== Completed Checklist Attached =====
    add_paragraph(doc, "Completed Checklist Attached:", bold=True, after_pt=2)
    _checkbox_line(doc, True, "Yes")
    _checkbox_line(doc, False, "No (include rationale if selected)", after_pt=14)

    # ===== Recommended Regulatory Action =====
    if assessment.get("manualReviewRequired"):
        submit_selected, action_note = False, "Manual Review Required (IVD — Flowchart D not supported)"
    else:
        submit_selected = bool(assessment["isSignificant"])
        action_note = "New 510(k)" if submit_selected else "Letter to file"

    add_paragraph(doc, "Recommended Regulatory Action:", bold=True, after_pt=2)
    _checkbox_line(doc, submit_selected, "Submit 510(k)")
    _checkbox_line(doc, not submit_selected, "Letter to file", after_pt=14)
    if assessment.get("manualReviewRequired"):
        add_paragraph(doc, action_note, italic=True, after_pt=14)

    # ===== Supporting Documents =====
    add_paragraph(doc, "Supporting Documents:", bold=True, after_pt=2)
    add_paragraph(doc, f"Design Specifications: {change_info.get('designSpecRef') or '[document number]'}", after_pt=2)
    add_paragraph(doc, f"Risk-Based Assessment: {change_info.get('riskAssessmentRef') or '[document number]'}", after_pt=2)
    add_paragraph(doc, f"Verification and Validation Summary: {change_info.get('vvSummaryRef') or '[document number]'}", after_pt=14)

    # ===== Signatures =====
    add_paragraph(doc, "Signatures", bold=True)
    add_signature_table(doc, [
        ("Prepared by", metadata.get("preparedBy", ""), metadata.get("preparedDate", "")),
        ("Reviewed by", metadata.get("reviewedBy", ""), metadata.get("reviewedDate", "")),
        ("Approved by", metadata.get("approvedBy", ""), metadata.get("approvedDate", "")),
    ])

    add_page_break(doc)

    # ===== Attachment: completed checklist =====
    add_paragraph(doc, "Attachment: Completed Checklist", bold=True, size=14)
    add_paragraph(
        doc,
        f'Assessed in accordance with FDA’s guidance "{FDA_GUIDANCE_TITLE}".',
        italic=True,
    )
    add_spacer(doc)

    add_paragraph(doc, "Main Flowchart Questions", bold=True)
    for q in assessment["path"]:
        if q["answer"] is None or not q["id"].startswith("MAIN"):
            continue
        _qa_checkboxes(doc, q["text_en"], q["answer"])

    last_chart_letter = None
    for q in assessment["path"]:
        if q["answer"] is None or q["id"].startswith("MAIN"):
            continue
        chart_letter = q["id"][0]
        if chart_letter != last_chart_letter:
            add_spacer(doc)
            title = CHART_TITLE.get(chart_letter, f"Chart {chart_letter}")
            add_paragraph(doc, title, bold=True)
            last_chart_letter = chart_letter
        _qa_checkboxes(doc, f"{q['id']}: {q['text_en']}", q["answer"])

    return doc
